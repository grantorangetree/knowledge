#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF → Gemma 4 Vision → 문서 변환기
pymupdf로 PDF 페이지를 이미지로 변환 후 Ollama로 분석
출력 형식: Markdown (.md) 또는 Word (.docx)
"""

import base64, json, sys, time, re
from pathlib import Path
from urllib import request, error
import fitz  # pymupdf
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── 설정 ──────────────────────────────────────────────
OLLAMA_URL    = "http://localhost:11434/api/generate"
DEFAULT_MODEL = "gemma4:e4b"
BASE_DIR      = Path(__file__).parent
OUTPUT_DIR    = BASE_DIR / "extracted_docs"

RESET  = "\033[0m";  BOLD   = "\033[1m"
CYAN   = "\033[96m"; GREEN  = "\033[92m"
YELLOW = "\033[93m"; RED    = "\033[91m"
PURPLE = "\033[95m"

def ok(m):   print(f"{GREEN}  ✓  {m}{RESET}")
def warn(m): print(f"{YELLOW}  ⚠  {m}{RESET}")
def err(m):  print(f"{RED}  ✗  {m}{RESET}")
def info(m): print(f"{CYAN}  →  {m}{RESET}")
def step(m): print(f"\n{BOLD}{CYAN}▸ {m}{RESET}")

PROMPT = """이 이미지는 스캔된 문서 또는 카탈로그의 한 페이지입니다. 다음을 수행해주세요:

1. 보이는 모든 텍스트를 정확하게 추출하세요
2. 제목, 소제목, 단락, 표, 목록 등 문서 구조를 유지하세요
3. 한국어와 영어가 섞여 있다면 그대로 보존하세요
4. 제품명, 스펙, 가격, 수치 등 중요한 정보는 특히 정확하게 추출하세요
5. 이미지/그림은 [이미지: 간단한 설명] 형식으로 표시하세요
6. 읽기 어려운 부분은 [불명확] 으로 표시하세요

추출된 내용만 출력하고, 불필요한 설명은 하지 마세요."""


def page_to_base64(page, dpi=150) -> str:
    """PDF 페이지를 base64 이미지로 변환"""
    mat = fitz.Matrix(dpi/72, dpi/72)
    pix = page.get_pixmap(matrix=mat)
    img_bytes = pix.tobytes("jpeg", jpg_quality=85)
    return base64.b64encode(img_bytes).decode("utf-8")


def analyze_page(b64_img: str, model: str, page_num: int) -> str:
    """Ollama API로 페이지 분석"""
    payload = {
        "model": model,
        "prompt": PROMPT,
        "images": [b64_img],
        "stream": False,
        "options": {
            "num_ctx": 8192,
            "temperature": 0.1,
        }
    }
    data = json.dumps(payload).encode("utf-8")
    req = request.Request(
        OLLAMA_URL,
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST"
    )
    try:
        with request.urlopen(req, timeout=180) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            return result.get("response", "").strip()
    except error.URLError as e:
        raise RuntimeError(f"Ollama 연결 실패: {e}")


# ── Word 문서 생성 ────────────────────────────────────
def markdown_to_docx(pages: dict, pdf_path: Path, model: str, total_pages: int) -> Path:
    """페이지별 텍스트를 Word 문서로 변환"""
    doc = Document()

    # 문서 제목 스타일 설정
    title = doc.add_heading(pdf_path.stem, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 메타 정보
    meta = doc.add_paragraph()
    meta.add_run(f"원본: {pdf_path.name}\n").italic = True
    meta.add_run(f"분석 모델: {model}\n").italic = True
    meta.add_run(f"총 페이지: {total_pages}").italic = True

    doc.add_page_break()

    for page_num in sorted(pages.keys()):
        text = pages[page_num]

        # 페이지 제목
        doc.add_heading(f"페이지 {page_num} / {total_pages}", level=1)

        # 텍스트를 라인 단위로 처리
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue

            # 마크다운 헤딩 처리
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            # 마크다운 표 처리 (| 로 시작하는 행)
            elif line.startswith("|") and line.endswith("|"):
                cells = [c.strip() for c in line.strip("|").split("|")]
                # 구분선은 건너뜀
                if all(set(c.replace("-","").replace(":","")) == set() or c.replace("-","").replace(":","") == "" for c in cells):
                    continue
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = "Table Grid"
                row = table.rows[0]
                for i, cell_text in enumerate(cells):
                    # 볼드 마크다운 제거
                    cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                    row.cells[i].text = cell_text
            # 목록 처리
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line[2:])
                p.add_run(clean)
            # 볼드 텍스트 처리
            elif re.search(r'\*\*(.+?)\*\*', line):
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            # 구분선
            elif line.startswith("---") or line.startswith("***"):
                doc.add_paragraph("─" * 40)
            # 일반 텍스트
            else:
                doc.add_paragraph(line)

        # 페이지 구분
        if page_num < total_pages:
            doc.add_page_break()

    out_path = OUTPUT_DIR / (pdf_path.stem + ".docx")
    doc.save(str(out_path))
    return out_path


def process_pdf(pdf_path: Path, model: str = DEFAULT_MODEL, start_page: int = 1, output_format: str = "both"):
    OUTPUT_DIR.mkdir(exist_ok=True)
    out_path = OUTPUT_DIR / (pdf_path.stem + ".md")

    print(f"""
{PURPLE}{BOLD}╔══════════════════════════════════════════════════════╗
║   PDF → 문서 변환기  (Gemma 4 Vision)                ║
╚══════════════════════════════════════════════════════╝{RESET}
""")

    step(f"PDF 로드: {pdf_path.name}")
    doc = fitz.open(str(pdf_path))
    total_pages = len(doc)
    ok(f"총 {total_pages}페이지")
    info(f"모델: {model}")
    info(f"출력: {out_path}")

    # 이어서 처리할 경우 기존 파일 확인
    existing_pages = set()
    if out_path.exists():
        content = out_path.read_text(encoding="utf-8")
        import re
        found = re.findall(r'## 페이지 (\d+)', content)
        existing_pages = set(int(p) for p in found)
        if existing_pages:
            warn(f"기존 처리된 페이지: {sorted(existing_pages)}")
            warn(f"이어서 처리합니다 (건너뜀: {len(existing_pages)}페이지)")

    results = {}
    failed_pages = []

    # 기존 내용 로드
    if out_path.exists() and existing_pages:
        raw = out_path.read_text(encoding="utf-8")
        # 기존 페이지별 내용 파싱
        sections = raw.split("\n## 페이지 ")
        for sec in sections[1:]:
            lines = sec.split("\n", 1)
            try:
                pnum = int(lines[0].split("/")[0].strip())
                results[pnum] = "## 페이지 " + sec
            except Exception:
                pass

    for page_num in range(1, total_pages + 1):
        if page_num < start_page:
            continue
        if page_num in existing_pages:
            info(f"페이지 {page_num}/{total_pages} — 건너뜀 (이미 처리됨)")
            continue

        print(f"\n  {BOLD}[{page_num}/{total_pages}]{RESET} 페이지 분석 중...", end="", flush=True)
        start = time.time()

        try:
            page = doc[page_num - 1]
            b64 = page_to_base64(page, dpi=150)
            text = analyze_page(b64, model, page_num)
            elapsed = time.time() - start

            if not text:
                text = "[내용 없음 또는 빈 페이지]"

            page_content = f"## 페이지 {page_num}/{total_pages}\n\n{text}\n"
            results[page_num] = page_content

            print(f" {GREEN}✓{RESET} ({elapsed:.1f}초)")

            # 매 페이지마다 중간 저장 (중단 시 복구 가능)
            sorted_results = [results[k] for k in sorted(results.keys())]
            header = f"# {pdf_path.name}\n\n> 원본: `{pdf_path}`  \n> 모델: `{model}`  \n> 총 페이지: {total_pages}\n\n---\n\n"
            out_path.write_text(header + "\n\n---\n\n".join(sorted_results), encoding="utf-8")

        except Exception as e:
            elapsed = time.time() - start
            print(f" {RED}✗{RESET} ({elapsed:.1f}초) — {e}")
            failed_pages.append(page_num)

        # API 과부하 방지
        if page_num < total_pages:
            time.sleep(0.5)

    doc.close()

    # Word 문서 생성
    word_path = None
    if output_format in ("docx", "both"):
        step("Word 문서 생성 중...")
        try:
            word_path = markdown_to_docx(results, pdf_path, model, total_pages)
            ok(f"Word 저장: {word_path.name}")
        except Exception as e:
            warn(f"Word 생성 실패: {e}")

    # 최종 결과
    print(f"\n{BOLD}{'═'*50}{RESET}")
    ok(f"완료! {len(results)}페이지 처리")
    if output_format in ("md", "both"):
        ok(f"Markdown: {out_path}")
    if word_path:
        ok(f"Word:     {word_path}")
    if failed_pages:
        warn(f"실패 페이지: {failed_pages}")
    if word_path:
        print(f"\n  {BOLD}📄 {word_path}{RESET}\n")
    else:
        print(f"\n  {BOLD}📄 {out_path}{RESET}\n")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("pdf", nargs="?", default="더마코-종합-E_-카달로그.pdf")
    parser.add_argument("--model", "-m", default=DEFAULT_MODEL)
    parser.add_argument("--start", "-s", type=int, default=1, help="시작 페이지")
    parser.add_argument("--format", "-f", choices=["md", "docx", "both"], default="both",
                        help="출력 형식: md / docx / both (기본: both)")
    args = parser.parse_args()

    pdf_path = Path(args.pdf)
    if not pdf_path.is_absolute():
        pdf_path = BASE_DIR / pdf_path
    if not pdf_path.exists():
        print(f"{RED}  ✗  파일 없음: {pdf_path}{RESET}")
        sys.exit(1)

    process_pdf(pdf_path, model=args.model, start_page=args.start, output_format=args.format)
